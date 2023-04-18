import docx


def refactor_docx(path, data):
    doc = docx.Document(path)

    # работа с обычным текстом (за исключением таблиц)
    for paragraph in doc.paragraphs:  # получаем данные всех параграфов
        for key in data.keys():  # сравниваем ключи и слова в параграфе
            if key in paragraph.text:  # если ключ находится параграфе, то заменяем его
                # задаём стиль при изменении слов
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = docx.shared.Pt(10)
                paragraph.text = paragraph.text.replace(key, data.get(key))

    # работа с таблицами
    for table in doc.tables:  # получаем все таблицы
        for row in table.rows:  # получаем все строчки таблицы
            for cell in row.cells:  # получаем все ячейки строки таблицы
                for paragraph in cell.paragraphs:  # получаем все параграфы в ячейке
                    for key in data.keys():  # сравниваем ключи с словами в параграфе
                        if key in paragraph.text:  # если ключ находится в параграфе, то меняем его
                            # задаём стиль при изменении слов
                            style = doc.styles['Normal']
                            font = style.font
                            font.name = 'Times New Roman'
                            font.size = docx.shared.Pt(10)
                            paragraph.text = paragraph.text.replace(key, data.get(key))

    doc.save(path)


# записываем в новый файл данные из словаря data
def write_docx(data, filename):
    doc = docx.Document()
    for key in data.keys():
        doc.add_paragraph(f'{key} - {data[key]}')
    doc.save(f'{filename}.docx')


if __name__ == '__main__':
    refactor_docx(path, data, filename)
    write_docx(data, filename)
